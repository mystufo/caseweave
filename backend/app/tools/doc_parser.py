"""Document parsing tool: extracts structured content from Word/PDF files.

Output contract (shared by both parsers)::

    {
        "chunks": [
            {"heading": str|None, "level": int, "content": str, "type": "text"|"table"},
            ...
        ],
        "tables": [{"table_index": int, "rows": list[list[str]]}, ...],
        "raw_text": str,           # 表格已序列化为 markdown，可直接喂 LLM
        "num_paragraphs"|"num_pages": int,
    }
"""
import io
from typing import Any, Iterable
from pathlib import Path


# ── 表格序列化 ────────────────────────────────────────────────────────────────

def _rows_to_markdown(rows: list[list[str]]) -> str:
    """把二维表格渲染成 markdown 表格字符串。
    空表返回空串；单行表也按 markdown 表头处理（保证 LLM 能识别为表）。
    """
    cleaned = [[(cell or "").replace("|", "\\|").replace("\n", " ").strip() for cell in row] for row in rows]
    cleaned = [r for r in cleaned if any(c for c in r)]
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    cleaned = [r + [""] * (width - len(r)) for r in cleaned]
    header = cleaned[0]
    body = cleaned[1:] if len(cleaned) > 1 else []
    sep = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _iter_docx_body(doc) -> Iterable[tuple[str, Any]]:
    """按文档原始顺序产出 ('paragraph', Paragraph) / ('table', Table)。
    python-docx 没有官方 API 做这件事，需要走 body XML 元素。
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    # 把 paragraphs/tables 按 XML 元素映射回 python-docx 对象
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}
    _ = (Paragraph, Table, para_iter, table_iter)  # 仅类型提示用

    for child in body.iterchildren():
        if child.tag == qn("w:p") and child in para_map:
            yield ("paragraph", para_map[child])
        elif child.tag == qn("w:tbl") and child in table_map:
            yield ("table", table_map[child])


def parse_docx(file_bytes: bytes) -> dict[str, Any]:
    """解析 .docx：按文档原顺序提取段落和表格，表格序列化为 markdown 并入 raw_text。"""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))

    chunks: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    raw_parts: list[str] = []

    current_heading: str | None = None
    current_level: int = 0
    text_buffer: list[str] = []
    table_counter = 0

    def flush_text():
        if text_buffer:
            content = "\n".join(text_buffer).strip()
            if content:
                chunks.append({
                    "heading": current_heading,
                    "level": current_level,
                    "content": content,
                    "type": "text",
                })
            text_buffer.clear()

    para_count = 0
    for kind, obj in _iter_docx_body(doc):
        if kind == "paragraph":
            para_count += 1
            text = obj.text.strip()
            if not text:
                continue
            style_name = obj.style.name if obj.style else ""
            if style_name.startswith("Heading"):
                flush_text()
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                current_heading = text
                current_level = level
                raw_parts.append(f"{'#' * min(level, 6)} {text}")
            else:
                text_buffer.append(text)
                raw_parts.append(text)
        else:  # table
            flush_text()
            rows = [[cell.text.strip() for cell in row.cells] for row in obj.rows]
            tables.append({"table_index": table_counter, "rows": rows})
            md = _rows_to_markdown(rows)
            if md:
                chunks.append({
                    "heading": current_heading,
                    "level": current_level,
                    "content": md,
                    "type": "table",
                })
                raw_parts.append(md)
            table_counter += 1

    flush_text()

    return {
        "chunks": chunks,
        "tables": tables,
        "raw_text": "\n\n".join(raw_parts),
        "num_paragraphs": para_count,
    }


# ── PDF ───────────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> dict[str, Any]:
    """解析 PDF：按页提取文本，并尝试抓表格序列化为 markdown 一并写入 raw_text。"""
    import pdfplumber

    chunks: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    raw_parts: list[str] = []
    table_counter = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_heading = f"Page {page_num}"
            raw_parts.append(f"## {page_heading}")

            text = (page.extract_text() or "").strip()
            if text:
                chunks.append({
                    "heading": page_heading,
                    "level": 1,
                    "content": text,
                    "type": "text",
                })
                raw_parts.append(text)

            # pdfplumber 自带表格抽取，可能误判，失败时静默跳过
            try:
                page_tables = page.extract_tables() or []
            except Exception:
                page_tables = []

            for raw_rows in page_tables:
                rows = [[(c or "").strip() for c in row] for row in raw_rows]
                if not any(any(c for c in r) for r in rows):
                    continue
                tables.append({"table_index": table_counter, "rows": rows})
                md = _rows_to_markdown(rows)
                if md:
                    chunks.append({
                        "heading": page_heading,
                        "level": 1,
                        "content": md,
                        "type": "table",
                    })
                    raw_parts.append(md)
                table_counter += 1

        num_pages = len(pdf.pages)

    return {
        "chunks": chunks,
        "tables": tables,
        "raw_text": "\n\n".join(raw_parts),
        "num_pages": num_pages,
    }


# ── 公共入口 ──────────────────────────────────────────────────────────────────

def parse_document(filename: str, file_bytes: bytes) -> dict[str, Any]:
    """Dispatch to the correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        return parse_docx(file_bytes)
    elif ext == ".pdf":
        return parse_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .docx, .pdf")


# ── 智能截断（给 LLM 用）─────────────────────────────────────────────────────

def get_doc_limit() -> int:
    """当前生效的文档字符上限，来自 .env 的 DOC_MAX_CHARS（默认 30000）。

    每次调用现取，不在 import 期固化——测试里 monkeypatch settings 才能生效。
    """
    from app.config import get_settings
    return get_settings().doc_max_chars


def truncate_for_llm(raw_text: str, limit: int | None = None) -> str:
    """超过 limit 时保留开头 + 末尾，中间用占位符提示被截断。
    需求文档的"验收标准/边界场景"经常在末尾，比单纯截前 N 字更靠谱。

    limit 缺省时取 DOC_MAX_CHARS；知识检索 query 那类场景由调用方显式传（如 limit=2000）。
    """
    if limit is None:
        limit = get_doc_limit()
    if len(raw_text) <= limit:
        return raw_text
    head_budget = int(limit * 0.7)
    tail_budget = limit - head_budget - 80  # 给中间提示留点 budget
    if tail_budget <= 0:
        return raw_text[:limit]
    head = raw_text[:head_budget].rstrip()
    tail = raw_text[-tail_budget:].lstrip()
    omitted = len(raw_text) - head_budget - tail_budget
    return (
        f"{head}\n\n"
        f"…（中间省略 {omitted} 字，原文共 {len(raw_text)} 字，请基于开头与末尾的内容做澄清/生成）…\n\n"
        f"{tail}"
    )


def doc_stats(*, chunks: int, tables: int, raw_text_length: int) -> dict[str, Any]:
    """上传/暂存接口回给前端的文档统计。

    `truncated` / `doc_char_limit` 在这里算好，前端直接用——不要在前端复制一份阈值常量，
    否则改了 DOC_MAX_CHARS 后端截断了、前端还不提示。
    """
    limit = get_doc_limit()
    return {
        "chunks": chunks,
        "tables": tables,
        "raw_text_length": raw_text_length,
        "truncated": raw_text_length > limit,
        "doc_char_limit": limit,
    }
